"""Build CAPSTONE PROJECT REPORT.docx in the Indian-university capstone format.

Formatting rules:
  - A4 paper, 1-inch (2.54 cm) margins on every side, portrait.
  - Times New Roman 12 pt body, 1.5 line spacing, justified, 6 pt space-after,
    0.5 inch first-line indent on body paragraphs (off for headings/lists/
    tables/captions/code/title page).
  - Heading 1 16 pt bold, page-break-before, 18 pt space-after.
  - Heading 2 14 pt bold, 12 pt space-after.
  - Heading 3 12 pt bold-italic, 6 pt space-after.
  - Title page centred, no page number.
  - Two sections: section 1 = title + front matter (lower-roman page numbers,
    titlePg suppressed); section 2 = body (decimal page numbers, restart at 1).
  - Live PAGE / NUMPAGES fields in footer, top-right header on body pages.
  - Real Word TOC field (\\o "1-3") so the user can update with F9.
  - Tables: dark-blue (#1F4E79) header row with white bold text, alternating
    body rows #F2F2F2 / white, "Table N: Title" caption above in 11 pt bold.
  - Figure placeholders: bordered box, italic centred text, caption below in
    11 pt bold.
  - Code listings: Consolas 10 pt, single line spacing, #F4F4F4 shading,
    0.25 inch left/right indent.
  - References: numbered with hanging 0.5 inch indent at 11 pt.

Section content lives in scripts/content.py and is unchanged.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

import content as C


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "CAPSTONE PROJECT REPORT.docx"

# Palette
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE_HEX = "FFFFFF"
HEADER_FILL = "1F4E79"        # dark blue, table header row
ROW_ALT_FILL = "F2F2F2"       # alternating body rows
CODE_FILL = "F4F4F4"          # code listing shading
GREY_TEXT_HEX = "595959"      # running header text
BORDER_GREY = "BFBFBF"        # figure box border
BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"


# ---------------------------------------------------------------------------
# low-level OOXML helpers
# ---------------------------------------------------------------------------

def _set_run_font(run, name: str = BODY_FONT) -> None:
    """Pin a run to a font for ASCII, hAnsi, cs and eastAsia ranges."""
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _set_style_font(style, name: str) -> None:
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _shade(element_tcPr_or_pPr, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element_tcPr_or_pPr.append(shd)


def set_cell_shading(cell, fill_hex: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), fill_hex)


def set_paragraph_borders(paragraph, color_hex: str = BORDER_GREY,
                          size: int = 8, space: int = 4) -> None:
    """Add a 1pt grey box border around a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), str(space))
        b.set(qn("w:color"), color_hex)
        pBdr.append(b)
    pPr.append(pBdr)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_numpages_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "NUMPAGES \\* MERGEFORMAT"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc_field(doc) -> None:
    """Insert a Word TOC field that renders the live table of contents on F9."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click here and choose 'Update Field' to populate "
        "the Table of Contents."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_section_pgnum(section, *, start: int, fmt: str) -> None:
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pgNum = OxmlElement("w:pgNumType")
    pgNum.set(qn("w:start"), str(start))
    pgNum.set(qn("w:fmt"), fmt)
    sectPr.append(pgNum)


def set_titlepg(section) -> None:
    sectPr = section._sectPr
    if sectPr.find(qn("w:titlePg")) is None:
        sectPr.append(OxmlElement("w:titlePg"))


# ---------------------------------------------------------------------------
# style setup
# ---------------------------------------------------------------------------

def setup_styles(doc) -> None:
    styles = doc.styles

    # Normal: every body paragraph inherits from this
    normal = styles["Normal"]
    _set_style_font(normal, BODY_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0.5)

    # Heading 1: chapter
    h1 = styles["Heading 1"]
    _set_style_font(h1, BODY_FONT)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = BLACK
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True

    # Heading 2: section
    h2 = styles["Heading 2"]
    _set_style_font(h2, BODY_FONT)
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = BLACK
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.first_line_indent = Inches(0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.page_break_before = False
    h2.paragraph_format.keep_with_next = True

    # Heading 3: subsection / sub-caption
    h3 = styles["Heading 3"]
    _set_style_font(h3, BODY_FONT)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = BLACK
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Inches(0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.page_break_before = False
    h3.paragraph_format.keep_with_next = True

    # Caption: table / figure caption
    if "ReportCaption" in [s.name for s in styles]:
        cap = styles["ReportCaption"]
    else:
        cap = styles.add_style("ReportCaption", WD_STYLE_TYPE.PARAGRAPH)
    cap.base_style = styles["Normal"]
    _set_style_font(cap, BODY_FONT)
    cap.font.size = Pt(11)
    cap.font.bold = True
    cap.font.italic = False
    cap.font.color.rgb = BLACK
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.line_spacing = 1.15
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = True

    # Code listing
    if "CodeListing" in [s.name for s in styles]:
        code = styles["CodeListing"]
    else:
        code = styles.add_style("CodeListing", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = styles["Normal"]
    _set_style_font(code, MONO_FONT)
    code.font.size = Pt(10)
    code.font.color.rgb = BLACK
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code.paragraph_format.first_line_indent = Inches(0)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(2)

    # Reference: hanging-indent numbered entries
    if "RefEntry" in [s.name for s in styles]:
        ref = styles["RefEntry"]
    else:
        ref = styles.add_style("RefEntry", WD_STYLE_TYPE.PARAGRAPH)
    ref.base_style = styles["Normal"]
    _set_style_font(ref, BODY_FONT)
    ref.font.size = Pt(11)
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref.paragraph_format.left_indent = Inches(0.5)
    ref.paragraph_format.first_line_indent = Inches(-0.5)
    ref.paragraph_format.line_spacing = 1.15
    ref.paragraph_format.space_after = Pt(6)


def setup_page_geometry(section) -> None:
    # A4
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # 1 inch margins
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    # header / footer 1.27 cm from edge so the content margin stays clean
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)


# ---------------------------------------------------------------------------
# header / footer wiring
# ---------------------------------------------------------------------------

def configure_section1_headers_footers(section) -> None:
    """Title page (suppressed) + roman footer for the rest of the front matter."""
    section.different_first_page_header_footer = True

    # First page (title): blank header and footer
    fp_header = section.first_page_header
    fp_header.is_linked_to_previous = False
    for p in list(fp_header.paragraphs):
        p.text = ""
    fp_footer = section.first_page_footer
    fp_footer.is_linked_to_previous = False
    for p in list(fp_footer.paragraphs):
        p.text = ""

    # Default footer for the front matter: centred lower-roman page number
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    add_page_number_field(p)
    for run in p.runs:
        _set_run_font(run, BODY_FONT)
        run.font.size = Pt(11)


def configure_section2_headers_footers(section) -> None:
    """Body: top-right header label, centred 'Page X of Y' footer."""
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.first_line_indent = Inches(0)
    hr = hp.add_run("LocalPulse | Capstone Report")
    _set_run_font(hr, BODY_FONT)
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor.from_string(GREY_TEXT_HEX)
    hr.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Inches(0)
    pre = fp.add_run("Page ")
    _set_run_font(pre, BODY_FONT)
    pre.font.size = Pt(11)
    add_page_number_field(fp)
    mid = fp.add_run(" of ")
    _set_run_font(mid, BODY_FONT)
    mid.font.size = Pt(11)
    add_numpages_field(fp)
    for run in fp.runs:
        _set_run_font(run, BODY_FONT)
        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# content helpers
# ---------------------------------------------------------------------------

def add_paragraphs(doc, paragraphs, *, italic: bool = False, indent: bool = True,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: int = 12,
                   bold: bool = False) -> None:
    for text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = align
        if not indent:
            p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(text)
        _set_run_font(run, BODY_FONT)
        run.font.size = Pt(size)
        run.italic = italic
        run.bold = bold


def add_centered(doc, text: str, *, size: int = 12, bold: bool = False,
                 italic: bool = False, caps: bool = False, color=None,
                 space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text.upper() if caps else text)
    _set_run_font(run, BODY_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def add_blank(doc, lines: int = 1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)


def add_heading(doc, text: str, level: int = 1, *, page_break_before=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.first_line_indent = Inches(0)
    if page_break_before is not None:
        h.paragraph_format.page_break_before = page_break_before
    for run in h.runs:
        _set_run_font(run, BODY_FONT)
    return h


def add_bullets(doc, items) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        bullet = p.add_run("•  ")
        _set_run_font(bullet, BODY_FONT)
        bullet.font.size = Pt(12)
        bullet.bold = True
        body = p.add_run(item)
        _set_run_font(body, BODY_FONT)
        body.font.size = Pt(12)


def add_numbered(doc, items) -> None:
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        head = p.add_run(f"{i}.  ")
        _set_run_font(head, BODY_FONT)
        head.font.size = Pt(12)
        head.bold = True
        body = p.add_run(item)
        _set_run_font(body, BODY_FONT)
        body.font.size = Pt(12)


def add_definition_list(doc, mapping) -> None:
    for key, val in mapping.items():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        head = p.add_run(f"{key}. ")
        _set_run_font(head, BODY_FONT)
        head.font.size = Pt(12)
        head.bold = True
        body = p.add_run(val)
        _set_run_font(body, BODY_FONT)
        body.font.size = Pt(12)


def add_table_caption(doc, text: str) -> None:
    p = doc.add_paragraph(style="ReportCaption")
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT)
    run.font.size = Pt(11)
    run.bold = True


def add_figure_caption(doc, text: str) -> None:
    p = doc.add_paragraph(style="ReportCaption")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT)
    run.font.size = Pt(11)
    run.bold = True


def add_table(doc, rows, *, first_row_header: bool = True,
              widths_in: list[float] | None = None) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells
        is_header = r_idx == 0 and first_row_header
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            cell.text = ""  # clear default
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Inches(0)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.15
            para.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = para.add_run(str(val))
            _set_run_font(run, BODY_FONT)
            run.font.size = Pt(11)
            if is_header:
                set_cell_shading(cell, HEADER_FILL)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE_HEX)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, ROW_ALT_FILL)

    if widths_in:
        for row in table.rows:
            for c_idx, w in enumerate(widths_in):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(w)


def add_code_block(doc, code: str, *, caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph(style="ReportCaption")
        cap.paragraph_format.first_line_indent = Inches(0)
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cap.add_run(caption)
        _set_run_font(run, BODY_FONT)
        run.font.size = Pt(11)
        run.bold = True

    for line in code.splitlines() or [""]:
        p = doc.add_paragraph(style="CodeListing")
        p.paragraph_format.first_line_indent = Inches(0)
        set_paragraph_shading(p, CODE_FILL)
        run = p.add_run(line if line else " ")
        _set_run_font(run, MONO_FONT)
        run.font.size = Pt(10)
    # closing spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Inches(0)
    spacer.paragraph_format.space_after = Pt(6)


def add_figure_placeholder(doc, label: str, caption: str) -> None:
    """A bordered placeholder box with italic centred label, then caption below."""
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.first_line_indent = Inches(0)
    box.paragraph_format.space_before = Pt(8)
    box.paragraph_format.space_after = Pt(0)
    box.paragraph_format.line_spacing = 1.5
    set_paragraph_borders(box)
    # Pad the box vertically with non-breaking line breaks so the box height
    # is roughly 4 cm regardless of caption length.
    pad_top = box.add_run(" \n")
    _set_run_font(pad_top, BODY_FONT)
    pad_top.font.size = Pt(10)
    body = box.add_run(f"[{label}: {caption}]")
    _set_run_font(body, BODY_FONT)
    body.font.size = Pt(11)
    body.italic = True
    body.font.color.rgb = RGBColor.from_string(GREY_TEXT_HEX)
    pad_bottom = box.add_run("\n ")
    _set_run_font(pad_bottom, BODY_FONT)
    pad_bottom.font.size = Pt(10)
    add_figure_caption(doc, f"{label}: {caption}")


def add_image(doc, rel_path: str, sub_caption: str, *, width_in: float = 6.0) -> None:
    """Embed a picture centred on the page, with an italic sub-caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(ROOT / rel_path), width=Inches(width_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = True
    cr = cap.add_run(sub_caption)
    _set_run_font(cr, BODY_FONT)
    cr.font.size = Pt(10)
    cr.italic = True
    cr.font.color.rgb = RGBColor.from_string(GREY_TEXT_HEX)


# ---------------------------------------------------------------------------
# section assemblers
# ---------------------------------------------------------------------------

def add_title_page(doc) -> None:
    """Centred title page; no page number (handled by titlePg + blank first
    footer). The first paragraph in the document must NOT carry the H1
    page-break-before that would push everything to a second sheet, so the
    title page deliberately uses plain centred paragraphs."""

    add_blank(doc, 1)

    add_centered(doc, C.PROJECT["university"], size=16, bold=True, caps=True,
                 space_after=6)
    add_centered(doc, C.PROJECT["school"], size=14, bold=True, caps=True,
                 space_after=4)
    add_centered(doc, C.PROJECT["location"], size=12, bold=False,
                 space_after=6)

    add_blank(doc, 3)

    add_centered(doc, C.PROJECT["title"], size=20, bold=True, space_after=18)

    add_centered(doc,
                 "Synopsis submitted for the partial fulfilment of the degree of",
                 size=12, italic=True, space_after=6)
    add_centered(doc, "BACHELOR OF TECHNOLOGY (CSE)", size=14, bold=True,
                 space_after=4)
    add_centered(doc, "in Cloud Computing", size=12, space_after=6)

    add_blank(doc, 2)

    add_centered(doc, "Submitted by", size=12, space_after=4)
    add_centered(doc, C.PROJECT["author"], size=14, bold=True, space_after=2)
    add_centered(doc, f"Roll No. {C.PROJECT['roll']}", size=12, space_after=10)

    add_centered(doc, "Under the supervision of", size=12, space_after=4)
    add_centered(doc, C.PROJECT["mentor"], size=14, bold=True, space_after=2)
    add_centered(doc, "Capstone Mentor", size=11, italic=True, space_after=6)

    add_blank(doc, 2)

    add_centered(doc, C.PROJECT["year"], size=14, bold=True, space_after=4)
    add_centered(doc, C.PROJECT["location"], size=12, space_after=2)


def add_front_matter(doc) -> None:
    add_heading(doc, "Acknowledgement", level=1)
    add_paragraphs(doc, C.ACKNOWLEDGEMENT)

    add_heading(doc, "Abstract", level=1)
    add_paragraphs(doc, C.ABSTRACT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    head = p.add_run("Keywords: ")
    _set_run_font(head, BODY_FONT)
    head.font.size = Pt(12)
    head.bold = True
    body = p.add_run(
        "crisis management, disaster response, earth observation, satellite "
        "data fusion, sensor cross-validation, self-learning calibration, "
        "conformal prediction, ECDSA P-256 provenance, multilingual voice "
        "assistant, Gemini Flash-Lite, Node.js, Cloud Run, accessibility, "
        "WCAG 2.2, DPDP Act 2023, Aatmanirbhar Bharat."
    )
    _set_run_font(body, BODY_FONT)
    body.font.size = Pt(12)
    body.italic = True

    add_heading(doc, "Table of Contents", level=1)
    add_toc_field(doc)

    add_heading(doc, "List of Figures", level=1)
    rows = [("Label", "Caption", "Page")] + [tuple(f) for f in C.FIGURES]
    add_table(doc, rows, widths_in=[1.0, 4.5, 0.7])

    add_heading(doc, "List of Tables", level=1)
    rows = [("Label", "Caption", "Page")] + [tuple(t) for t in C.TABLES_LIST]
    add_table(doc, rows, widths_in=[1.0, 4.5, 0.7])


def add_body(doc) -> None:
    # ---- Chapter 1 ----------------------------------------------------------
    h = add_heading(doc, "Chapter 1. Introduction and Problem Definition",
                    level=1, page_break_before=False)
    add_heading(doc, "1.1 Background", level=2)
    add_paragraphs(doc, C.INTRO_BACKGROUND)
    add_heading(doc, "1.2 Problem Statement", level=2)
    add_paragraphs(doc, C.INTRO_PROBLEM_STATEMENT)
    add_heading(doc, "1.3 Objectives", level=2)
    add_numbered(doc, C.INTRO_OBJECTIVES)
    add_heading(doc, "1.4 Scope", level=2)
    add_paragraphs(doc, C.INTRO_SCOPE)
    add_heading(doc, "1.5 Target Users", level=2)
    add_bullets(doc, C.INTRO_TARGET_USERS)
    add_heading(doc, "1.6 Significance", level=2)
    add_paragraphs(doc, C.INTRO_SIGNIFICANCE)

    # ---- Chapter 2 ----------------------------------------------------------
    add_heading(doc, "Chapter 2. System Requirements", level=1)
    add_heading(doc, "2.1 Functional Requirements", level=2)
    add_bullets(doc, C.REQ_FUNCTIONAL)
    add_heading(doc, "2.2 Non-functional Requirements", level=2)
    add_definition_list(doc, C.REQ_NONFUNCTIONAL)
    add_heading(doc, "2.3 Hardware and Software Requirements", level=2)
    add_definition_list(doc, C.REQ_HARDWARE_SOFTWARE)

    # ---- Chapter 3 ----------------------------------------------------------
    add_heading(doc, "Chapter 3. System Architecture and Design", level=1)
    add_heading(doc, "3.1 Architectural Overview", level=2)
    add_paragraphs(doc, C.ARCH_OVERVIEW)
    add_figure_placeholder(
        doc, "Figure 1", "High-level System Architecture of LocalPulse"
    )

    add_heading(doc, "3.2 Component Inventory", level=2)
    rows = [("Component", "Description")] + list(C.ARCH_COMPONENTS)
    add_table(doc, rows, widths_in=[2.2, 4.0])

    add_heading(doc, "3.3 Data Flow", level=2)
    add_paragraphs(doc, C.ARCH_DATAFLOW)
    add_figure_placeholder(
        doc, "Figure 2",
        "End-to-end Data Flow from Public Source to Resident Device"
    )

    add_heading(doc, "3.4 API Design", level=2)
    add_paragraphs(doc, C.ARCH_API_DESIGN)

    add_heading(doc, "3.5 Data and State Design", level=2)
    add_paragraphs(doc, C.ARCH_DATA_DESIGN)

    add_heading(doc, "3.6 Security Architecture", level=2)
    add_paragraphs(doc, C.ARCH_SECURITY)
    add_table_caption(doc, "Table 4: STRIDE Threat Model and Mitigations")
    rows = [
        ("Threat (STRIDE)", "Vector", "Mitigation"),
        ("Spoofing", "Fake resident reports",
         "Agentic verifier drops contradicted reports; cross-source corroboration gates trust."),
        ("Tampering", "Man-in-the-middle; forged hazard warning",
         "TLS 1.3, HSTS preload, SRI; ECDSA P-256 hash-chained provenance on every EO warning."),
        ("Repudiation", "Disputed or backdated warning",
         "Structured audit logs with correlation IDs; tamper-evident, offline-verifiable certificate chain."),
        ("Information Disclosure", "Leak of personal data",
         "Report form collects no name, email or phone; no PII in logs or URLs."),
        ("Denial of Service", "Flood of requests",
         "Cloudflare edge absorbs floods; hot read path served from memory, no DB round-trip."),
        ("Elevation of Privilege", "Compromised account; budget abuse",
         "Least-privilege Cloud Run SA, keyless WIF; /tasks/ingest gated by constant-time token."),
    ]
    add_table(doc, rows, widths_in=[1.6, 2.1, 2.5])

    add_heading(doc, "3.7 Deployment Topology", level=2)
    add_paragraphs(doc, C.ARCH_DEPLOYMENT)

    # ---- Chapter 4 ----------------------------------------------------------
    add_heading(doc, "Chapter 4. Technology Stack", level=1)
    add_paragraphs(doc, [
        "Table 1 lists every technology choice in the LocalPulse stack with "
        "the reason it was selected over alternatives. The selection is "
        "guided by three constraints: idle cost close to zero, mobile-first "
        "performance on a 3G connection and accessibility at WCAG 2.2 level "
        "AAA."
    ])
    add_table_caption(doc, "Table 1: Technology Choices and Rationale")
    add_table(doc, C.TECH_STACK, widths_in=[1.4, 2.0, 3.0])

    # ---- Chapter 5 ----------------------------------------------------------
    add_heading(doc, "Chapter 5. Implementation", level=1)
    add_heading(doc, "5.1 Code Organisation", level=2)
    add_paragraphs(doc, C.IMPL_CODE_ORG)

    add_heading(doc, "5.2 Public API Endpoints", level=2)
    add_table_caption(doc, "Table 2: Public HTTP and JSON API Endpoints")
    rows = [("Endpoint", "Description")] + list(C.IMPL_KEY_APIS)
    add_table(doc, rows, widths_in=[2.4, 3.8])

    add_heading(doc, "5.3 Internationalisation", level=2)
    add_paragraphs(doc, C.IMPL_I18N)

    add_heading(doc, "5.4 Mobile-First Responsive Layout", level=2)
    add_paragraphs(doc, C.IMPL_RESPONSIVE)
    add_figure_placeholder(
        doc, "Figure 4", "Resident Dashboard Wireframe (mobile, 360x800)"
    )
    add_figure_placeholder(
        doc, "Figure 5", "Responder Console Wireframe (tablet, 1024x768)"
    )

    # Figure 8: real product-UI screenshots, stacked with sub-captions.
    add_image(doc, "public/img/ui-dashboard-light.jpg", "(a) Resident dashboard, light theme")
    add_image(doc, "public/img/ui-dashboard-dark.jpg", "(b) Resident dashboard, dark theme")
    add_image(doc, "public/img/ui-satellite-light.jpg", "(c) Satellite Intelligence panel, light theme")
    add_image(doc, "public/img/ui-satellite-dark.jpg", "(d) Satellite Intelligence panel, dark theme")
    add_figure_caption(
        doc,
        "Figure 8: Deployed User Interface at localpulse.dmj.one. The same "
        "responsive, theme-aware layout is shown in light and dark themes: "
        "(a, b) the resident dashboard and (c, d) the redesigned Satellite "
        "Intelligence panel."
    )

    add_heading(doc, "5.5 Real-time Update Loop", level=2)
    add_paragraphs(doc, C.IMPL_REALTIME)

    add_heading(doc, "5.6 AI Summarisation Pipeline", level=2)
    add_paragraphs(doc, C.IMPL_AI_PIPELINE)

    add_heading(doc, "5.7 Voice Bot Flow", level=2)
    add_paragraphs(doc, C.IMPL_VOICE_FLOW)
    add_figure_placeholder(
        doc, "Figure 3", "In-Browser Voice Flow over the Web Speech API, Grounded in Live Data"
    )

    add_heading(doc, "5.8 Selected Code Listings", level=2)
    for snip in C.IMPL_CODE_SNIPPETS:
        add_code_block(doc, snip["code"], caption=snip["caption"])

    # ---- Chapter 6 ----------------------------------------------------------
    add_heading(doc, "Chapter 6. Algorithms and Models", level=1)
    add_heading(doc, "6.1 Social-media Summarisation Pipeline", level=2)
    add_paragraphs(doc, C.ALGO_SUMMARISATION)
    add_heading(doc, "6.2 Voice Intent Classification", level=2)
    add_paragraphs(doc, C.ALGO_INTENT)
    add_heading(doc, "6.3 Trust Score", level=2)
    add_paragraphs(doc, C.ALGO_TRUST)
    add_heading(doc, "6.4 Language Identification", level=2)
    add_paragraphs(doc, C.ALGO_LANGID)

    add_heading(doc, "6.5 Earth-Observation Subsystem", level=2)
    add_paragraphs(doc, C.EO_OVERVIEW)
    add_figure_placeholder(
        doc, "Figure 7",
        "Earth-Observation Fusion, Cross-Validation and Signed Provenance"
    )
    add_heading(doc, "6.5.1 Multi-sensor Satellite Fusion", level=3)
    add_paragraphs(doc, C.EO_FUSION)
    add_table_caption(doc, "Table 7: Earth-Observation Sensor Adapters and Hazard Axes")
    add_table(doc, list(C.EO_ADAPTERS), widths_in=[1.7, 1.6, 1.3, 1.6])
    add_heading(doc, "6.5.2 Cross-validation and Anti-spoofing", level=3)
    add_paragraphs(doc, C.EO_DIVERGENCE)
    add_heading(doc, "6.5.3 Forecasting and the Self-learning World Engine", level=3)
    add_paragraphs(doc, C.EO_FORECAST_WORLD)
    add_heading(doc, "6.5.4 Tamper-evident, Offline-verifiable Provenance", level=3)
    add_paragraphs(doc, C.EO_PROVENANCE)
    add_heading(doc, "6.5.5 Forensic Warning Certificate and Route Clearance", level=3)
    add_paragraphs(doc, C.EO_CERTIFICATE_ROUTE)

    # ---- Chapter 7 ----------------------------------------------------------
    add_heading(doc, "Chapter 7. Testing", level=1)
    add_paragraphs(doc, [
        "The automated suite uses the built-in Node test runner, node:test, "
        "invoked with node --test: one hundred and two cases, all passing, in "
        "roughly fourteen seconds, with no third-party test dependency to "
        "install or audit. The cases concentrate on the hardest and most novel "
        "subsystem, earth observation, across the focus areas below."
    ])
    add_table_caption(doc, "Table 5: Test Suite Coverage and Sample Results")
    rows = [("Focus area", "Runner", "Coverage")] + list(C.TEST_LAYERS)
    add_table(doc, rows, widths_in=[1.4, 1.2, 3.8])
    add_heading(doc, "7.1 Sample Test Results", level=2)
    add_table(doc, C.TEST_SAMPLE_RESULTS, widths_in=[2.0, 1.2, 1.2, 1.6])

    # ---- Chapter 8 ----------------------------------------------------------
    add_heading(doc, "Chapter 8. Results and Performance Analysis", level=1)
    add_paragraphs(doc, [
        "The minimum lovable product was measured against a set of "
        "engineering targets that map back to the non-functional "
        "requirements in Section 2.2. Table 3 reports the target value, the "
        "achieved value and a short note on the measurement context."
    ])
    add_table_caption(doc, "Table 3: Performance Targets and Achieved Numbers")
    add_table(doc, C.RESULTS_TARGETS, widths_in=[1.9, 1.2, 1.2, 2.1])
    add_heading(doc, "8.1 Discussion", level=2)
    add_paragraphs(doc, C.RESULTS_DISCUSSION)

    # ---- Chapter 9 ----------------------------------------------------------
    add_heading(doc, "Chapter 9. Deployment", level=1)
    add_paragraphs(doc, C.DEPLOY_OVERVIEW)
    add_figure_placeholder(
        doc, "Figure 6", "Cloud Run Deployment Topology in asia-east1"
    )
    add_heading(doc, "9.1 Dockerfile", level=2)
    add_code_block(doc, C.DEPLOY_DOCKERFILE,
                   caption="Listing 4: Multi-stage Dockerfile.")
    add_heading(doc, "9.2 Cloud Run Deploy", level=2)
    add_code_block(doc, C.DEPLOY_GCLOUD,
                   caption="Listing 5: gcloud run deploy with custom domain mapping.")
    add_heading(doc, "9.3 Continuous Integration and Release", level=2)
    add_paragraphs(doc, C.DEPLOY_CICD)

    # ---- Chapter 10 ---------------------------------------------------------
    add_heading(doc, "Chapter 10. Challenges and Solutions", level=1)
    for ch in C.CHALLENGES:
        add_heading(doc, ch["title"], level=2)
        for label, key in (("Problem.", "problem"),
                           ("Solution.", "solution"),
                           ("Outcome.", "outcome")):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0)
            head = p.add_run(f"{label} ")
            _set_run_font(head, BODY_FONT)
            head.font.size = Pt(12)
            head.bold = True
            body = p.add_run(ch[key])
            _set_run_font(body, BODY_FONT)
            body.font.size = Pt(12)

    # ---- Chapter 11 ---------------------------------------------------------
    add_heading(doc, "Chapter 11. Conclusion and Future Scope", level=1)
    add_heading(doc, "11.1 Conclusion", level=2)
    add_paragraphs(doc, C.CONCLUSION)
    add_heading(doc, "11.2 Future Scope", level=2)
    add_bullets(doc, C.FUTURE_SCOPE)
    add_paragraphs(doc, C.CONCLUSION_CLOSE, italic=True)

    # ---- Viva ---------------------------------------------------------------
    add_heading(doc, "Questions and Answers (Viva)", level=1)
    add_paragraphs(doc, [
        "The following ten questions are the standard viva voce set for this "
        "capstone. Each answer is written to be defensible on the stand."
    ])
    for i, qa in enumerate(C.VIVA_QA, 1):
        # Numbered Q&A: bold question prefix + question, then answer paragraph.
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        qp.paragraph_format.first_line_indent = Inches(0)
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        qp.paragraph_format.keep_with_next = True
        q_run = qp.add_run(f"Q{i}. {qa['q']}")
        _set_run_font(q_run, BODY_FONT)
        q_run.font.size = Pt(12)
        q_run.bold = True

        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ap.paragraph_format.first_line_indent = Inches(0)
        ap.paragraph_format.left_indent = Inches(0.25)
        ans_label = ap.add_run("Answer. ")
        _set_run_font(ans_label, BODY_FONT)
        ans_label.font.size = Pt(12)
        ans_label.bold = True
        ans = ap.add_run(qa["a"])
        _set_run_font(ans, BODY_FONT)
        ans.font.size = Pt(12)

    # ---- References ---------------------------------------------------------
    add_heading(doc, "References", level=1)
    for i, ref in enumerate(C.REFERENCES, 1):
        p = doc.add_paragraph(style="RefEntry")
        p.paragraph_format.first_line_indent = Inches(-0.5)
        idx_run = p.add_run(f"[{i}]  ")
        _set_run_font(idx_run, BODY_FONT)
        idx_run.font.size = Pt(11)
        idx_run.bold = True
        body_run = p.add_run(ref)
        _set_run_font(body_run, BODY_FONT)
        body_run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# top-level build
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    setup_styles(doc)

    # ---- Section 1: title page + front matter (lower-roman page numbers) ----
    section1 = doc.sections[0]
    setup_page_geometry(section1)
    set_section_pgnum(section1, start=0, fmt="lowerRoman")
    set_titlepg(section1)
    configure_section1_headers_footers(section1)

    add_title_page(doc)
    add_front_matter(doc)

    # ---- Section 2: body (decimal page numbers, restart at 1) --------------
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page_geometry(section2)
    set_section_pgnum(section2, start=1, fmt="decimal")
    section2.different_first_page_header_footer = False
    configure_section2_headers_footers(section2)

    add_body(doc)

    doc.save(str(DOCX))
    size = DOCX.stat().st_size
    print(f"Wrote {DOCX} ({size} bytes, {size / 1024:.1f} KiB)")


if __name__ == "__main__":
    build()
